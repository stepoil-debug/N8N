from __future__ import annotations

import csv
import email
import hmac
import io
import json
import os
import re
import tempfile
import uuid
from datetime import UTC, datetime
from email import policy
from pathlib import Path
from typing import Any

from docx import Document
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

app = FastAPI(title="STEP Industrial Audit Document API", version="1.0.0")
ARTIFACT_ROOT = Path(os.getenv("ARTIFACT_ROOT", "/data/artifacts"))
ARTIFACT_ROOT.mkdir(parents=True, exist_ok=True)
MAX_UPLOAD = int(os.getenv("MAX_UPLOAD_MB", "100")) * 1024 * 1024
SAFE = re.compile(r"[^A-Za-z0-9._-]+")


def auth(x_step_api_key: str | None = Header(default=None)) -> None:
    expected = os.getenv("API_KEY", "")
    if expected and (not x_step_api_key or not hmac.compare_digest(expected, x_step_api_key)):
        raise HTTPException(status_code=401, detail="API key inválida")


def safe(value: str, fallback: str = "item") -> str:
    result = SAFE.sub("-", value.strip()).strip(".-")
    return result[:120] or fallback


def workspace(opportunity_id: str) -> Path:
    path = ARTIFACT_ROOT / safe(opportunity_id, "opportunity")
    path.mkdir(parents=True, exist_ok=True)
    return path


def artifact(opportunity_id: str, path: Path) -> dict[str, Any]:
    return {
        "opportunity_id": opportunity_id,
        "artifact_name": path.name,
        "artifact_path": str(path),
        "download_path": f"/v1/artifacts/{safe(opportunity_id)}/{path.name}",
        "size_bytes": path.stat().st_size,
        "created_at": datetime.now(UTC).isoformat(),
    }


def pdf_content(data: bytes) -> tuple[list[dict[str, Any]], list[str]]:
    pages = [{"page": i, "text": (page.extract_text() or "").strip()} for i, page in enumerate(PdfReader(io.BytesIO(data)).pages, 1)]
    warnings: list[str] = []
    if not any(page["text"] for page in pages):
        warnings.append("PDF sem camada de texto; OCR aplicado.")
        try:
            import pytesseract
            from pdf2image import convert_from_bytes
            pages = [{"page": i, "text": pytesseract.image_to_string(image, lang="por+eng").strip()} for i, image in enumerate(convert_from_bytes(data, dpi=200), 1)]
        except Exception as exc:
            warnings.append(f"OCR indisponível: {exc}")
    return pages, warnings


def extract(filename: str, data: bytes) -> dict[str, Any]:
    suffix = Path(filename).suffix.casefold()
    warnings: list[str] = []
    if suffix == ".pdf":
        content, warnings = pdf_content(data)
        kind = "pdf"
    elif suffix in {".xlsx", ".xlsm"}:
        book = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
        content = [{"sheet": ws.title, "rows": [["" if value is None else str(value) for value in row] for row in ws.iter_rows(values_only=True) if any(value not in (None, "") for value in row)]} for ws in book.worksheets]
        kind = "spreadsheet"
    elif suffix == ".docx":
        doc = Document(io.BytesIO(data))
        content = [{"paragraphs": [p.text for p in doc.paragraphs if p.text.strip()], "tables": [[[cell.text for cell in row.cells] for row in table.rows] for table in doc.tables]}]
        kind = "document"
    elif suffix == ".eml":
        msg = email.message_from_bytes(data, policy=policy.default)
        bodies = []
        attachments = []
        for part in msg.walk():
            if part.get_content_disposition() == "attachment":
                attachments.append(part.get_filename() or "anexo")
            elif part.get_content_type() == "text/plain":
                try: bodies.append(part.get_content())
                except Exception: pass
        content = [{"from": str(msg.get("from", "")), "to": str(msg.get("to", "")), "subject": str(msg.get("subject", "")), "date": str(msg.get("date", "")), "body": "\n\n".join(bodies), "attachments": attachments}]
        kind = "email"
    elif suffix == ".msg":
        import extract_msg
        with tempfile.NamedTemporaryFile(suffix=".msg") as handle:
            handle.write(data); handle.flush(); msg = extract_msg.Message(handle.name)
            content = [{"from": msg.sender or "", "to": msg.to or "", "cc": msg.cc or "", "subject": msg.subject or "", "date": str(msg.date or ""), "body": msg.body or "", "attachments": [getattr(a, "longFilename", None) or getattr(a, "shortFilename", None) or "anexo" for a in msg.attachments]}]
            msg.close()
        kind = "email"
    elif suffix == ".csv":
        content = [{"rows": list(csv.reader(io.StringIO(data.decode("utf-8-sig", errors="replace"))))}]
        kind = "spreadsheet"
    elif suffix in {".txt", ".md", ".json"}:
        text = data.decode("utf-8-sig", errors="replace")
        content = [{"text": text}]
        kind = "text"
    else:
        raise HTTPException(status_code=415, detail=f"Extensão não suportada: {suffix or 'sem extensão'}")
    return {"filename": Path(filename).name, "extension": suffix, "kind": kind, "content": content, "warnings": warnings}


@app.get("/health")
def health() -> dict[str, str]:
    return {"status": "ok", "service": "document-api", "version": app.version}


@app.post("/v1/opportunities/prepare", dependencies=[Depends(auth)])
async def prepare(request: Request) -> dict[str, Any]:
    body = await request.json()
    oid = str(body.get("opportunity_id", "")).strip()
    if not oid: raise HTTPException(status_code=422, detail="opportunity_id é obrigatório")
    meta = {"opportunity_id": oid, "client": body.get("client"), "rfq_id": body.get("rfq_id"), "created_at": datetime.now(UTC).isoformat()}
    (workspace(oid) / "opportunity.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"status": "prepared", **meta}


@app.post("/v1/documents/extract", dependencies=[Depends(auth)])
async def extract_document(file: UploadFile = File(...), opportunity_id: str = Form(default="unassigned")) -> dict[str, Any]:
    data = await file.read()
    if len(data) > MAX_UPLOAD: raise HTTPException(status_code=413, detail="Arquivo acima do limite configurado")
    return extract(file.filename or "documento", data) | {"opportunity_id": opportunity_id, "size_bytes": len(data)}


@app.post("/v1/triage/checklist", dependencies=[Depends(auth)])
async def checklist(request: Request) -> dict[str, Any]:
    body = await request.json(); oid = str(body.get("opportunity_id", "")).strip(); analysis = body.get("analysis")
    if not oid or not isinstance(analysis, dict): raise HTTPException(status_code=422, detail="opportunity_id e analysis são obrigatórios")
    requirements = analysis.get("requirements") or []
    if not isinstance(requirements, list): raise HTTPException(status_code=422, detail="requirements deve ser lista")
    output = workspace(oid) / safe(str(body.get("output_name") or f"Checklist_RFQ_{oid}.xlsx"))
    if output.suffix.casefold() != ".xlsx": output = output.with_suffix(".xlsx")
    wb = Workbook(); ws = wb.active; ws.title = "Checklist"
    headers = ["ID", "Categoria", "Requisito", "Origem / Documento", "Localização da evidência", "Evidência resumida", "Tipo sugerido pela IA", "Prioridade sugerida", "Tratamento sugerido", "Decisão do usuário", "Tipo final", "Prioridade final", "Tratamento final", "Responsável final"]
    ws.append(headers)
    navy = PatternFill("solid", fgColor="0B2D4D")
    for cell in ws[1]: cell.fill = navy; cell.font = Font(color="FFFFFF", bold=True); cell.alignment = Alignment(wrap_text=True, horizontal="center")
    for index, item in enumerate(requirements, 1):
        ws.append([item.get("id") or f"RFQ-{index:03d}", item.get("category", ""), item.get("requirement", ""), item.get("source_document", ""), item.get("source_location", ""), item.get("source_evidence", ""), item.get("type_ai", ""), item.get("priority_ai", ""), item.get("treatment_ai", ""), item.get("decision", "A confirmar"), item.get("type_final", ""), item.get("priority_final", ""), item.get("treatment_final", ""), item.get("owner_final", "")])
    widths = [16, 24, 65, 30, 28, 50, 23, 22, 34, 22, 20, 20, 30, 24]
    for i, width in enumerate(widths, 1): ws.column_dimensions[chr(64+i) if i <= 26 else "A"].width = width
    ws.freeze_panes = "A2"; ws.auto_filter.ref = ws.dimensions
    summary = wb.create_sheet("Resumo"); summary.append(["Cliente", (analysis.get("summary") or {}).get("client", "")]); summary.append(["RFQ", (analysis.get("summary") or {}).get("rfq_id", "")]); summary.append(["Total de requisitos", len(requirements)])
    wb.save(output)
    return artifact(oid, output)


def inventory_from_file(path: Path) -> dict[str, Any]:
    wb = load_workbook(path, data_only=True); ws = wb["Checklist"]
    headers = {str(cell.value).strip().casefold(): cell.column for cell in ws[1] if cell.value}
    def col(*names: str) -> int | None:
        for name in names:
            if name.casefold() in headers: return headers[name.casefold()]
        return None
    idc, reqc = col("ID"), col("Requisito")
    if not idc or not reqc: raise HTTPException(status_code=422, detail="Checklist sem colunas ID/Requisito")
    rows = []
    for row in range(2, ws.max_row + 1):
        rid = str(ws.cell(row, idc).value or "").strip(); requirement = str(ws.cell(row, reqc).value or "").strip()
        if not rid and not requirement: continue
        decision = str(ws.cell(row, col("Decisão do usuário") or 1).value or "").strip()
        final_type = str(ws.cell(row, col("Tipo final") or 1).value or "").strip()
        rows.append({"checklist_id": rid or f"ROW-{row:04d}", "source_sheet": ws.title, "source_row": row, "requirement": requirement, "category": str(ws.cell(row, col("Categoria") or 1).value or "").strip(), "requirement_source": str(ws.cell(row, col("Origem / Documento") or 1).value or "").strip(), "requirement_location": str(ws.cell(row, col("Localização da evidência") or 1).value or "").strip(), "requirement_evidence": str(ws.cell(row, col("Evidência resumida") or 1).value or "").strip(), "applicable": decision.casefold() != "excluir" and final_type.casefold() != "não aplicável"})
    return {"source_file": path.name, "source_sheet": ws.title, "total_nonempty_rows": len(rows), "applicable_rows": sum(1 for row in rows if row["applicable"]), "checklist_inventory": rows}


@app.post("/v1/adherence/inventory", dependencies=[Depends(auth)])
async def inventory(file: UploadFile = File(...), opportunity_id: str = Form(default="unassigned")) -> dict[str, Any]:
    data = await file.read()
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as handle: handle.write(data); source = Path(handle.name)
    try: return inventory_from_file(source) | {"opportunity_id": opportunity_id}
    finally: source.unlink(missing_ok=True)


@app.post("/v1/adherence/inventory/from-artifact", dependencies=[Depends(auth)])
async def inventory_artifact(request: Request) -> dict[str, Any]:
    body = await request.json(); oid = str(body.get("opportunity_id", "")).strip(); name = safe(str(body.get("artifact_name", "")))
    path = workspace(oid) / name
    if not path.is_file(): raise HTTPException(status_code=404, detail="Checklist não encontrado")
    return inventory_from_file(path) | {"opportunity_id": oid}


@app.post("/v1/adherence/report", dependencies=[Depends(auth)])
async def report(request: Request) -> dict[str, Any]:
    body = await request.json(); oid = str(body.get("opportunity_id", "")).strip(); analysis = body.get("analysis")
    if not oid or not isinstance(analysis, dict): raise HTTPException(status_code=422, detail="opportunity_id e analysis são obrigatórios")
    inventory = [item for item in analysis.get("checklist_inventory", []) if item.get("applicable", True)]
    items = analysis.get("items") or []
    inventory_ids = {str(item.get("checklist_id", "")) for item in inventory if item.get("checklist_id")}
    mapped_ids = {str(item.get("checklist_id") or item.get("id") or "") for item in items}
    unmapped = sorted(inventory_ids - mapped_ids)
    no_points = [str(item.get("checklist_id") or item.get("id") or "") for item in items if not item.get("comparison_points")]
    if (unmapped or no_points) and not body.get("allow_incomplete"):
        raise HTTPException(status_code=422, detail={"message": "Cobertura incompleta", "unmapped": unmapped, "items_without_comparison_points": no_points})
    scores = {"Atendido": 1.0, "Parcial": 0.5, "Não atendido": 0.0}; weights = {"Crítico": 4, "Alto": 3, "Médio": 2, "Baixo": 1, "Informativo": 0}
    num = den = 0.0; blockers = []; rows = [["ID", "Critério", "Pedido", "Proposta", "Status"]]
    for item in items:
        for point in item.get("comparison_points") or []:
            ass = point.get("human_decision") or point.get("ai_assessment") or {}; status = ass.get("status", "Não verificável"); impact = ass.get("impact", "Médio"); weight = float(weights.get(impact, 0))
            if status in scores and weight > 0: num += scores[status] * weight; den += weight
            if ass.get("blocks") and status in {"Não atendido", "Parcial"}: blockers.append(point.get("id"))
            rows.append([point.get("id", ""), point.get("parameter", ""), point.get("requested", ""), point.get("proposed", ""), status])
    adherence = round(num / den * 100, 1) if den else None; coverage = round(len(inventory_ids & mapped_ids) / len(inventory_ids) * 100, 1) if inventory_ids else None
    decision = "ANÁLISE INCOMPLETA" if unmapped or no_points else "NÃO ACEITÁVEL" if blockers else "ACEITÁVEL" if adherence is not None and adherence >= 95 else "ACEITÁVEL COM RESSALVAS" if adherence is not None and adherence >= 85 else "NÃO ACEITÁVEL"
    output = workspace(oid) / safe(str(body.get("output_name") or f"Validacao_Aderencia_{oid}.pdf"))
    if output.suffix.casefold() != ".pdf": output = output.with_suffix(".pdf")
    styles = getSampleStyleSheet(); doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=12*mm, leftMargin=12*mm, topMargin=12*mm, bottomMargin=12*mm)
    story = [Paragraph("VALIDAÇÃO DE ADERÊNCIA DA PROPOSTA", styles["Title"]), Spacer(1, 5*mm), Paragraph(f"Decisão: <b>{decision}</b>", styles["Heading2"]), Paragraph(f"Aderência: {adherence if adherence is not None else '-'}% | Cobertura: {coverage if coverage is not None else '-'}%", styles["BodyText"]), Spacer(1, 5*mm)]
    table = Table(rows, repeatRows=1, colWidths=[25*mm, 45*mm, 42*mm, 42*mm, 30*mm]); table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#0B2D4D")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("GRID", (0,0), (-1,-1), 0.35, colors.grey), ("VALIGN", (0,0), (-1,-1), "TOP"), ("FONTSIZE", (0,0), (-1,-1), 7)])); story.append(table); doc.build(story)
    result = artifact(oid, output); result.update({"decision": decision, "adherence_percent": adherence, "coverage_percent": coverage, "blocking_points": blockers}); return result


@app.get("/v1/artifacts/{opportunity_id}/{artifact_name}", dependencies=[Depends(auth)])
def download(opportunity_id: str, artifact_name: str) -> FileResponse:
    path = workspace(opportunity_id) / safe(artifact_name)
    if not path.is_file(): raise HTTPException(status_code=404, detail="Artefato não encontrado")
    return FileResponse(path, filename=path.name)
