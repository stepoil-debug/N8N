from __future__ import annotations

import json
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from .main import artifact, safe, workspace

SEVERITY_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "informational": 4}
SEVERITY_PT = {
    "critical": "Crítico",
    "high": "Alto",
    "medium": "Médio",
    "low": "Baixo",
    "informational": "Informativo",
}
STATUS_PT = {
    "met": "Atendido",
    "partial": "Parcial",
    "not_met": "Não atendido",
    "not_verifiable": "Não verificável",
}


def _list(value: Any) -> list[dict[str, Any]]:
    return [item for item in (value or []) if isinstance(item, dict)]


def _text(value: Any, fallback: str = "") -> str:
    if value is None:
        return fallback
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value).strip()


def _analysis_summary(analysis: dict[str, Any]) -> dict[str, Any]:
    findings = _list(analysis.get("findings"))
    requirements = _list(analysis.get("requirements"))
    commitments = _list(analysis.get("commitments"))
    counts = {key: 0 for key in SEVERITY_ORDER}
    for finding in findings:
        severity = _text(finding.get("severity"), "medium").casefold()
        if severity not in counts:
            severity = "medium"
        counts[severity] += 1
    statuses = {"met": 0, "partial": 0, "not_met": 0, "not_verifiable": 0}
    for requirement in requirements:
        status = _text(requirement.get("status"), "not_verifiable").casefold()
        if status in statuses:
            statuses[status] += 1
    covered = statuses["met"] + statuses["partial"] + statuses["not_met"]
    coverage = round(covered / len(requirements) * 100, 1) if requirements else None
    weighted = statuses["met"] + statuses["partial"] * 0.5
    adherence = round(weighted / covered * 100, 1) if covered else None
    blocking = [f for f in findings if bool(f.get("blocking")) or _text(f.get("severity")).casefold() == "critical"]
    supplied = analysis.get("summary") if isinstance(analysis.get("summary"), dict) else {}
    recommendation = _text(supplied.get("recommendation") or analysis.get("recommendation"), "review_before_submit")
    risk_level = _text(supplied.get("risk_level"), "critical" if blocking else "high" if counts["high"] else "medium")
    return {
        "recommendation": recommendation,
        "risk_level": risk_level,
        "requirements_total": len(requirements),
        "commitments_total": len(commitments),
        "findings_total": len(findings),
        "blocking_risks": len(blocking),
        "severity_counts": counts,
        "status_counts": statuses,
        "coverage_percent": supplied.get("coverage_percent", coverage),
        "adherence_percent": supplied.get("adherence_percent", adherence),
        "executive_opinion": _text(supplied.get("executive_opinion") or analysis.get("executive_opinion")),
    }


def normalize_analysis(analysis: dict[str, Any], opportunity: dict[str, Any], package: dict[str, Any]) -> dict[str, Any]:
    normalized = {
        "opportunity": opportunity,
        "package": {
            "package_name": package.get("package_name"),
            "root_folder": package.get("root_folder"),
            "summary": package.get("summary", {}),
        },
        "requirements": _list(analysis.get("requirements")),
        "commitments": _list(analysis.get("commitments")),
        "findings": _list(analysis.get("findings")),
        "corrections": _list(analysis.get("corrections")),
        "corrected_proposal": analysis.get("corrected_proposal") if isinstance(analysis.get("corrected_proposal"), dict) else {},
        "assumptions": [str(item) for item in (analysis.get("assumptions") or [])],
        "not_verifiable": _list(analysis.get("not_verifiable")),
        "generated_at": datetime.now(UTC).isoformat(),
        "model": _text(analysis.get("model")),
    }
    normalized["findings"].sort(key=lambda item: SEVERITY_ORDER.get(_text(item.get("severity"), "medium").casefold(), 2))
    normalized["summary"] = _analysis_summary(normalized | {"summary": analysis.get("summary", {})})
    return normalized


def _cell(value: Any) -> str:
    return _text(value)


def _write_excel(path: Path, data: dict[str, Any]) -> None:
    wb = Workbook()
    navy = PatternFill("solid", fgColor="0B2D4D")
    white_bold = Font(color="FFFFFF", bold=True)

    def sheet(name: str, headers: list[str], rows: list[list[Any]], widths: list[int]) -> None:
        ws = wb.create_sheet(name) if wb.worksheets else wb.active
        ws.title = name
        ws.append(headers)
        for cell in ws[1]:
            cell.fill = navy
            cell.font = white_bold
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for row in rows:
            ws.append([_cell(value) for value in row])
        for idx, width in enumerate(widths, 1):
            ws.column_dimensions[ws.cell(1, idx).column_letter].width = width
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        ws.freeze_panes = "A2"
        ws.auto_filter.ref = ws.dimensions

    summary = data["summary"]
    sheet("Resumo", ["Indicador", "Valor"], [
        ["Recomendação", summary.get("recommendation")],
        ["Nível de risco", summary.get("risk_level")],
        ["Parecer executivo", summary.get("executive_opinion")],
        ["Requisitos", summary.get("requirements_total")],
        ["Compromissos", summary.get("commitments_total")],
        ["Achados", summary.get("findings_total")],
        ["Bloqueios", summary.get("blocking_risks")],
        ["Cobertura (%)", summary.get("coverage_percent")],
        ["Aderência (%)", summary.get("adherence_percent")],
    ], [32, 110])
    sheet("Requisitos", ["ID", "Categoria", "Requisito", "Status", "Documento", "Localização", "Evidência"], [
        [item.get("id"), item.get("category"), item.get("requirement") or item.get("description"), STATUS_PT.get(_text(item.get("status")).casefold(), item.get("status")), item.get("source_document"), item.get("source_location"), item.get("source_evidence")]
        for item in data["requirements"]
    ], [16, 22, 62, 18, 34, 28, 60])
    sheet("Compromissos STEP", ["ID", "Categoria", "Compromisso", "Documento", "Localização", "Evidência"], [
        [item.get("id"), item.get("category"), item.get("commitment") or item.get("description"), item.get("source_document"), item.get("source_location"), item.get("source_evidence")]
        for item in data["commitments"]
    ], [16, 22, 68, 34, 28, 60])
    sheet("Achados", ["ID", "Severidade", "Categoria", "Inconsistência", "Impacto", "Evidência cliente", "Evidência STEP", "Correção necessária", "Bloqueante"], [
        [item.get("id"), SEVERITY_PT.get(_text(item.get("severity")).casefold(), item.get("severity")), item.get("category"), item.get("title") or item.get("inconsistency"), item.get("impact"), item.get("client_evidence"), item.get("step_evidence"), item.get("required_correction") or item.get("recommendation"), "SIM" if item.get("blocking") else "NÃO"]
        for item in data["findings"]
    ], [15, 16, 22, 55, 42, 58, 58, 62, 14])
    sheet("Correções", ["ID", "Seção da proposta", "Texto atual", "Texto corrigido", "Motivo", "Validação humana"], [
        [item.get("id"), item.get("section"), item.get("current_text"), item.get("corrected_text"), item.get("reason"), "SIM" if item.get("requires_human_validation") else "NÃO"]
        for item in data["corrections"]
    ], [15, 28, 60, 72, 55, 18])
    wb.save(path)


def _proposal_sections(data: dict[str, Any]) -> list[dict[str, Any]]:
    proposal = data.get("corrected_proposal") or {}
    sections = _list(proposal.get("sections"))
    if sections:
        return sections
    corrections = data.get("corrections") or []
    grouped: dict[str, list[str]] = {}
    for item in corrections:
        title = _text(item.get("section"), "Correções recomendadas")
        grouped.setdefault(title, []).append(_text(item.get("corrected_text") or item.get("required_correction") or item.get("reason")))
    return [{"title": title, "paragraphs": values} for title, values in grouped.items()]


def _write_corrected_docx(path: Path, data: dict[str, Any]) -> None:
    opportunity = data.get("opportunity") or {}
    proposal = data.get("corrected_proposal") or {}
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PROPOSTA TÉCNICO-COMERCIAL REVISADA")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{_text(opportunity.get('opportunity_id'))} · {_text(opportunity.get('client'))}").bold = True
    if opportunity.get("rfq_id"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"RFQ: {_text(opportunity.get('rfq_id'))}")

    warning = doc.add_paragraph()
    warning.style = doc.styles["Intense Quote"]
    warning.add_run("Documento gerado a partir da auditoria automatizada. Valores, prazos e obrigações classificados como não verificáveis ou sujeitos a aprovação devem ser confirmados pelo responsável antes do envio ao cliente.")

    intro = _text(proposal.get("introduction"))
    if intro:
        doc.add_heading("1. Apresentação", level=1)
        doc.add_paragraph(intro)

    for index, section_data in enumerate(_proposal_sections(data), 2):
        doc.add_heading(f"{index}. {_text(section_data.get('title'), 'Seção')}", level=1)
        for paragraph in section_data.get("paragraphs") or []:
            if isinstance(paragraph, dict):
                text = _text(paragraph.get("text") or paragraph.get("content"))
                if paragraph.get("bullet"):
                    doc.add_paragraph(text, style="List Bullet")
                else:
                    doc.add_paragraph(text)
            else:
                doc.add_paragraph(_text(paragraph))
        for item in section_data.get("bullets") or []:
            doc.add_paragraph(_text(item), style="List Bullet")
        table_data = section_data.get("table")
        if isinstance(table_data, dict):
            headers = [str(h) for h in table_data.get("headers") or []]
            rows = table_data.get("rows") or []
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for idx, header in enumerate(headers):
                    table.rows[0].cells[idx].text = header
                for row in rows:
                    cells = table.add_row().cells
                    for idx, value in enumerate(row[:len(headers)]):
                        cells[idx].text = _text(value)

    exclusions = proposal.get("exclusions") or []
    if exclusions:
        doc.add_heading("Exclusões e ressalvas", level=1)
        for item in exclusions:
            doc.add_paragraph(_text(item), style="List Bullet")

    validations = [item for item in data.get("corrections") or [] if item.get("requires_human_validation")]
    if validations:
        doc.add_heading("Itens pendentes de validação humana", level=1)
        for item in validations:
            doc.add_paragraph(f"{_text(item.get('section'))}: {_text(item.get('corrected_text') or item.get('reason'))}", style="List Bullet")

    doc.add_paragraph()
    footer = doc.add_paragraph("STEP Oil & Gas · Revisão gerada em " + datetime.now().strftime("%d/%m/%Y %H:%M"))
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
    doc.save(path)


def _p(value: Any, style: Any) -> Paragraph:
    return Paragraph(escape(_text(value)).replace("\n", "<br/>"), style)


def _write_audit_pdf(path: Path, data: dict[str, Any]) -> None:
    styles = getSampleStyleSheet()
    small = ParagraphStyle("small", parent=styles["BodyText"], fontSize=7, leading=9)
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=12 * mm, leftMargin=12 * mm, topMargin=12 * mm, bottomMargin=12 * mm)
    summary = data["summary"]
    story: list[Any] = [Paragraph("AUDITORIA DE ADERÊNCIA DA PROPOSTA", styles["Title"]), Spacer(1, 4 * mm)]
    story.append(_p(f"Oportunidade: {(data.get('opportunity') or {}).get('opportunity_id', '')} · Cliente: {(data.get('opportunity') or {}).get('client', '')}", styles["Heading3"]))
    story.append(_p(f"Recomendação: {summary.get('recommendation')} | Risco: {summary.get('risk_level')} | Bloqueios: {summary.get('blocking_risks')}", styles["Heading2"]))
    story.append(_p(summary.get("executive_opinion") or "Parecer executivo não informado pelo agente.", styles["BodyText"]))
    story.append(Spacer(1, 4 * mm))
    rows: list[list[Any]] = [["ID", "Severidade", "Inconsistência", "Impacto", "Correção"]]
    for item in data["findings"]:
        rows.append([
            _p(item.get("id"), small),
            _p(SEVERITY_PT.get(_text(item.get("severity")).casefold(), item.get("severity")), small),
            _p(item.get("title") or item.get("inconsistency"), small),
            _p(item.get("impact"), small),
            _p(item.get("required_correction") or item.get("recommendation"), small),
        ])
    table = Table(rows, repeatRows=1, colWidths=[18 * mm, 20 * mm, 50 * mm, 45 * mm, 50 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0B2D4D")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
    ]))
    story.append(table)
    story.append(PageBreak())
    story.append(Paragraph("Correções propostas", styles["Heading1"]))
    for item in data["corrections"]:
        story.append(_p(f"{item.get('id', '')} · {item.get('section', '')}", styles["Heading3"]))
        story.append(_p(item.get("corrected_text") or item.get("reason"), styles["BodyText"]))
    doc.build(story)


def _write_proposal_pdf(path: Path, data: dict[str, Any]) -> None:
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm)
    opportunity = data.get("opportunity") or {}
    proposal = data.get("corrected_proposal") or {}
    story: list[Any] = [
        Paragraph("PROPOSTA TÉCNICO-COMERCIAL REVISADA", styles["Title"]),
        Spacer(1, 3 * mm),
        _p(f"{opportunity.get('opportunity_id', '')} · {opportunity.get('client', '')} · RFQ {opportunity.get('rfq_id', '')}", styles["Heading3"]),
        Spacer(1, 5 * mm),
        _p("Documento gerado pela auditoria automatizada. Itens não verificáveis devem ser validados antes da submissão.", styles["Italic"]),
    ]
    if proposal.get("introduction"):
        story.extend([Paragraph("Apresentação", styles["Heading1"]), _p(proposal.get("introduction"), styles["BodyText"])])
    for section in _proposal_sections(data):
        story.append(Paragraph(escape(_text(section.get("title"), "Seção")), styles["Heading1"]))
        for paragraph in section.get("paragraphs") or []:
            story.append(_p(paragraph.get("text") if isinstance(paragraph, dict) else paragraph, styles["BodyText"]))
        for item in section.get("bullets") or []:
            story.append(_p("• " + _text(item), styles["BodyText"]))
    doc.build(story)


def generate_artifacts(opportunity: dict[str, Any], package: dict[str, Any], analysis: dict[str, Any]) -> dict[str, Any]:
    oid = _text(opportunity.get("opportunity_id"), "opportunity")
    data = normalize_analysis(analysis, opportunity, package)
    root = workspace(oid)
    json_path = root / safe(f"Auditoria_Completa_{oid}.json")
    xlsx_path = root / safe(f"Auditoria_Completa_{oid}.xlsx")
    audit_pdf = root / safe(f"Relatorio_Auditoria_{oid}.pdf")
    proposal_docx = root / safe(f"Proposta_Revisada_{oid}.docx")
    proposal_pdf = root / safe(f"Proposta_Revisada_{oid}.pdf")
    json_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    _write_excel(xlsx_path, data)
    _write_audit_pdf(audit_pdf, data)
    _write_corrected_docx(proposal_docx, data)
    _write_proposal_pdf(proposal_pdf, data)
    artifacts = [artifact(oid, path) for path in (audit_pdf, proposal_docx, proposal_pdf, xlsx_path, json_path)]
    return {
        "status": "analysis_completed",
        "opportunity": opportunity,
        "summary": data["summary"],
        "findings": data["findings"],
        "corrections": data["corrections"],
        "requirements": data["requirements"],
        "commitments": data["commitments"],
        "artifacts": artifacts,
        "completed_at": datetime.now(UTC).isoformat(),
    }
