from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK

from .main import safe, workspace


def _text(value: Any) -> str:
    return "" if value is None else str(value).strip()


def _replace_paragraph(paragraph: Any, old: str, new: str) -> bool:
    if not old or old.casefold() in {"not_verifiable", "não verificável"}:
        return False
    full = "".join(run.text for run in paragraph.runs)
    if old not in full:
        return False
    replaced = full.replace(old, new)
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)
    return True


def _all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def revise_original_proposal(opportunity: dict[str, Any], analysis: dict[str, Any]) -> Path | None:
    oid = _text(opportunity.get("opportunity_id")) or "opportunity"
    root = workspace(oid)
    candidates = sorted(root.glob("Proposta_Original_*.docx"), key=lambda path: path.stat().st_mtime, reverse=True)
    if not candidates:
        return None

    source = candidates[0]
    output = root / safe(f"Proposta_Revisada_{oid}.docx")
    document = Document(source)
    corrections = [item for item in (analysis.get("corrections") or []) if isinstance(item, dict)]
    applied: list[dict[str, Any]] = []
    pending: list[dict[str, Any]] = []

    for correction in corrections:
        old = _text(correction.get("current_text"))
        new = _text(correction.get("corrected_text"))
        changed = False
        if old and new:
            for paragraph in _all_paragraphs(document):
                if _replace_paragraph(paragraph, old, new):
                    changed = True
        target = applied if changed else pending
        target.append(correction)

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading("ANEXO — REVISÃO DA AUDITORIA", level=1)
    document.add_paragraph(
        "Este anexo registra as correções recomendadas pela auditoria automatizada. "
        "Itens marcados para validação humana devem ser confirmados antes da emissão ao cliente."
    )

    if applied:
        document.add_heading("Correções aplicadas diretamente", level=2)
        for item in applied:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{_text(item.get('section'))}: ").bold = True
            paragraph.add_run(_text(item.get("corrected_text")))

    if pending:
        document.add_heading("Correções incluídas para validação", level=2)
        for item in pending:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{_text(item.get('section'))}: ").bold = True
            paragraph.add_run(_text(item.get("corrected_text") or item.get("reason")))
            if item.get("requires_human_validation"):
                paragraph.add_run(" [VALIDAÇÃO HUMANA OBRIGATÓRIA]").bold = True

    document.add_paragraph(f"Revisão gerada em {datetime.now().strftime('%d/%m/%Y %H:%M')}.")
    document.save(output)
    return output
